from flask import Flask, request, jsonify
from flask_cors import CORS
import spacy
import requests
from docx import Document
import os

app = Flask(__name__)
CORS(app)

# Lade deutsches NER-Modell
try:
    nlp = spacy.load("de_core_news_lg")
except:
    raise RuntimeError("Bitte fÃ¼hre aus: python -m spacy download de_core_news_sm")

# DOCX-Inhalt auslesen (auch Tabellen)
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = []
    text.extend(para.text for para in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    return "\n".join(text)

# Orte erkennen
def extract_places(text):
    doc = nlp(text)
    return list(set(ent.text for ent in doc.ents if ent.label_ in ("LOC", "GPE", "FAC")))

# Gazetteer-Suche
def search_gazetteer(place):
    url = f"https://gazetteer.dainst.org/search.json?q={place}"
    try:
        resp = requests.get(url)
        resp.raise_for_status()
        return resp.json().get("result", [])
    except Exception:
        return []

# API-Endpunkt fÃ¼r Datei-Upload
@app.route("/api/analyse", methods=["POST"])
def analyse():
    if "file" not in request.files:
        return jsonify({"error": "Keine Datei empfangen."}), 400

    file = request.files["file"]
    if not file.filename.endswith(".docx"):
        return jsonify({"error": "Nur .docx-Dateien erlaubt."}), 400

    temp_path = os.path.join("temp.docx")
    file.save(temp_path)

    text = extract_text_from_docx(temp_path)
    places = extract_places(text)
    results = []

    for place in places:
        gaz_hits = search_gazetteer(place)
        simplified = []

        if gaz_hits:
            for res in gaz_hits[:5]:
                simplified.append({
                    "place_query": place,
                    "title": res.get("prefName", {}).get("title", "-"),
                    "uri": res.get("uri") or f"https://gazetteer.dainst.org/place/{res.get('gazId', '')}",
                    "coords": res.get("prefLocation", {}).get("coordinates", []),
                    "types": res.get("types", [])
                })

        # Auch ohne Treffer aufnehmen
        results.append({
            "query": place,
            "results": simplified
        })

    os.remove(temp_path)
    return jsonify(results)


if __name__ == "__main__":
    app.run(debug=True)
